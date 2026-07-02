from pathlib import Path

import fitz
from docx import Document

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".wav", ".mp3", ".txt"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
AUDIO_EXTENSIONS = {".wav", ".mp3"}


def detect_file_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in DOCUMENT_EXTENSIONS:
        return "document"
    if suffix in IMAGE_EXTENSIONS:
        return "image"
    if suffix in AUDIO_EXTENSIONS:
        return "audio"
    raise ValueError(f"Unsupported file extension: {suffix}")


def extract_pdf_text(path: Path) -> str:
    with fitz.open(path) as doc:
        return "\n".join(page.get_text("text") for page in doc)


def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    table_cells = [
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    return "\n".join([*paragraphs, *table_cells])


def extract_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def clean_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.replace("\x00", " ").splitlines()]
    return "\n".join(line for line in lines if line).strip()


def extract_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return clean_text(extract_pdf_text(path))
    if suffix == ".docx":
        return clean_text(extract_docx_text(path))
    if suffix == ".txt":
        return clean_text(extract_text_file(path))
    raise ValueError(f"Unsupported document extension: {suffix}")
